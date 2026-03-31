from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
def set_font(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
    return p

def para(text="", bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold, size=size)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(11)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F497D")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    p.paragraph_format.left_indent = Inches(0.3)

# ── Title Page ───────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI-Based Phishing Detection System")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Comprehensive Project Report")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
doc.add_page_break()

# ── 1. Overview ──────────────────────────────────────────────────────────────
heading("1. Project Overview")
para(
    "This project is a hybrid, multi-signal phishing detection system that combines Machine Learning, "
    "real-time HTML analysis, and external threat intelligence APIs to classify URLs as phishing, "
    "suspicious, or benign. It is designed as a production-ready web service with a visual UI, "
    "REST API, and automated alerting via n8n."
)

# ── 2. Objectives ────────────────────────────────────────────────────────────
heading("2. Objectives")
for obj in [
    "Detect phishing URLs in real time using multiple independent signals.",
    "Provide a transparent, explainable risk score with human-readable reasons.",
    "Automate alerting to Slack and log all threats via n8n workflows.",
    "Be deployable to the cloud (Render) with minimal configuration.",
]:
    bullet(obj)

# ── 3. Architecture ──────────────────────────────────────────────────────────
heading("3. System Architecture")
code_block(
    "User\n"
    " └─► Streamlit UI (ui/app.py)\n"
    "       └─► FastAPI Backend (api/main.py)\n"
    "             ├─► ML Model          (ml_model/predict.py)\n"
    "             ├─► HTML Scraper      (scraper/html_scraper.py)\n"
    "             ├─► VirusTotal API    (threat_intel/api_checker.py)\n"
    "             ├─► AbuseIPDB API     (threat_intel/api_checker.py)\n"
    "             ├─► IPStack API       (threat_intel/api_checker.py)\n"
    "             └─► Decision Engine   (decision/engine.py)\n"
    "                   └─► Final Score + Label + Reasons\n\n"
    "n8n Automation\n"
    " └─► Webhook Trigger (POST /check-url)\n"
    "       └─► FastAPI /webhook\n"
    "             ├─► If phishing  → Log Alert (/alert) → Slack\n"
    "             └─► If suspicious → Slack Warning"
)

# ── 4. Module Breakdown ───────────────────────────────────────────────────────
heading("4. Module Breakdown")

heading("4.1  API Layer — api/main.py", level=2)
para("Built with FastAPI v2.0.0. Key features:")
for f in [
    "SSRF Protection: _is_safe_url() blocks private/loopback/link-local IPs before any external call.",
    "TTL Cache: Results cached for 300 seconds per URL to reduce redundant API calls.",
    "URL Normalization: Automatically prepends http:// if scheme is missing.",
    "Async Parallelism: HTML scraping and all three threat intel APIs run concurrently via asyncio.gather().",
]:
    bullet(f)
para("API Endpoints:", bold=True)
add_table(
    ["Method", "Endpoint", "Description"],
    [
        ["GET", "/", "Health check"],
        ["GET", "/health", "Cache size + status"],
        ["POST", "/predict", "Main URL analysis"],
        ["GET/POST", "/webhook", "n8n-compatible endpoint"],
        ["POST", "/alert", "Receive alert from n8n"],
        ["GET", "/alerts", "View all logged alerts"],
        ["DELETE", "/cache", "Clear result cache"],
    ],
)

heading("4.2  Machine Learning — ml_model/", level=2)
para("Training (train.py):", bold=True)
for f in [
    "Dataset: Kaggle Web Page Phishing Detection Dataset (dataset.csv).",
    "Algorithm: XGBoost Classifier — n_estimators=400, max_depth=7, learning_rate=0.05.",
    "Regularization: reg_alpha=0.1, reg_lambda=1.0, gamma=0.1.",
    "Validation: 5-fold Stratified Cross-Validation (F1-weighted).",
    "Output: phishing_model.pkl (model + feature list + label encoder).",
]:
    bullet(f)
para("Prediction (predict.py):", bold=True)
para(
    "Extracts 87 URL and page features covering URL structure (length, dots, hyphens, special chars, "
    "subdomains, TLD), content signals (phishing keywords, brand impersonation, shortening services), "
    "page-level features (hyperlink ratios, login forms, iframes, popups, favicons, media ratios), "
    "and DNS resolution. Returns label, confidence (%), and per-class probabilities."
)

heading("4.3  HTML Scraper — scraper/html_scraper.py", level=2)
para("Fetches and parses the target page using BeautifulSoup + lxml. Extracted features:")
add_table(
    ["Feature", "Description"],
    [
        ["has_login_form", "Presence of any <form> tag"],
        ["has_password_field", "Presence of <input type='password'>"],
        ["form_action_empty", "Form with empty or '#' action attribute"],
        ["external_script_count", "Scripts loaded from external domains"],
        ["iframe_count", "Number of <iframe> tags"],
        ["hidden_element_count", "Elements with display:none style"],
        ["link_count", "Total anchor tags"],
        ["image_count", "Total image tags"],
        ["has_favicon", "Presence of favicon link tag"],
    ],
)

heading("4.4  Threat Intelligence — threat_intel/api_checker.py", level=2)
para("All three APIs are called asynchronously in parallel via check_all().")
para("VirusTotal:", bold=True)
para("Submits URL for analysis and polls the result. Returns malicious, suspicious, harmless, and undetected engine counts.")
para("AbuseIPDB:", bold=True)
para("Resolves URL to IP and checks abuse history over the last 90 days. Returns abuse_score (%), total_reports, and country.")
para("IPStack:", bold=True)
para("Provides geolocation and threat intelligence per IP. Returns country, city, coordinates, is_tor, is_proxy, is_anonymous, and is_attacker flags.")

heading("4.5  Decision Engine — decision/engine.py", level=2)
para("Aggregates all signals into a single 0–100 risk score:")
add_table(
    ["Source", "Max Points", "Trigger Condition"],
    [
        ["ML Model", "40", "Phishing label × confidence ratio"],
        ["VirusTotal", "35", "Malicious engine ratio or suspicious count"],
        ["AbuseIPDB", "12", "Abuse score > 20% or > 10 reports"],
        ["HTML Analysis", "10", "Login form + password field, empty action, iframes"],
        ["IPStack", "13", "Attacker (13), TOR/Proxy (9), Anonymous (5)"],
    ],
)
para("Label Thresholds:", bold=True)
for t in [
    "score >= 70  →  phishing",
    "score >= 40  →  suspicious",
    "score < 40   →  benign",
]:
    bullet(t)

heading("4.6  UI — ui/app.py", level=2)
para("Built with Streamlit and Plotly. Features:")
for f in [
    "URL input with single-click analysis.",
    "Color-coded result header: red (phishing), yellow (suspicious), green (benign).",
    "Risk score metric + progress bar.",
    "Plotly gauge chart for visual risk level.",
    "Bar chart for per-source score contribution.",
    "Expandable sections: ML details, VirusTotal, AbuseIPDB, IPStack (with world map), HTML features.",
    "Cache indicator and response time display.",
]:
    bullet(f)

heading("4.7  n8n Automation — n8n/phishing_detection_workflow.json", level=2)
para("Automated workflow with 8 nodes:")
code_block(
    "Webhook Trigger (POST /check-url)\n"
    "  └─► Call Phishing API (POST /webhook)\n"
    "        ├─► Is Phishing?   → Log Alert (/alert) → Slack Alert\n"
    "        ├─► Is Suspicious? → Slack Warning\n"
    "        └─► Respond to Webhook"
)
for f in [
    "Webhook URL: http://localhost:5678/webhook/check-url",
    "Slack messages include URL, score, and reasons.",
    "All phishing alerts are persisted via POST /alert.",
]:
    bullet(f)

# ── 5. Tech Stack ─────────────────────────────────────────────────────────────
heading("5. Tech Stack")
add_table(
    ["Layer", "Technology"],
    [
        ["UI", "Streamlit, Plotly"],
        ["API", "FastAPI, Uvicorn"],
        ["ML", "XGBoost, scikit-learn, joblib"],
        ["HTML Parsing", "BeautifulSoup4, lxml"],
        ["HTTP Clients", "requests (sync), httpx (async)"],
        ["Threat Intel", "VirusTotal, AbuseIPDB, IPStack"],
        ["Automation", "n8n"],
        ["Config", "python-dotenv"],
        ["Deployment", "Render (render.yaml)"],
    ],
)

# ── 6. Security ───────────────────────────────────────────────────────────────
heading("6. Security Considerations")
add_table(
    ["Concern", "Mitigation"],
    [
        ["SSRF", "_is_safe_url() blocks private/loopback/reserved IPs before any request"],
        ["API Key Exposure", "Keys loaded from .env, never hardcoded; .gitignore excludes .env"],
        ["Input Validation", "Pydantic models validate all request bodies; URL normalization applied"],
        ["Cache Poisoning", "TTL-based cache (5 min) limits stale data window"],
        ["Denial of Service", "Async parallel calls with timeouts (6–30s per service)"],
    ],
)

# ── 7. Deployment ─────────────────────────────────────────────────────────────
heading("7. Deployment")
heading("7.1  Local", level=2)
code_block(
    "pip install -r requirements.txt\n"
    "uvicorn api.main:app --reload        # Terminal 1\n"
    "streamlit run ui/app.py              # Terminal 2"
)
heading("7.2  Cloud (Render)", level=2)
for f in [
    "Defined in render.yaml — two services: API + UI.",
    "Environment variables set in Render dashboard: VIRUSTOTAL_API_KEY, ABUSEIPDB_API_KEY, API_URL.",
    "Model file phishing_model.pkl must be committed before deploy.",
]:
    bullet(f)

# ── 8. End-to-End Data Flow ───────────────────────────────────────────────────
heading("8. End-to-End Data Flow")
steps = [
    "User submits URL via Streamlit UI.",
    "UI sends POST /predict to FastAPI.",
    "API checks TTL cache → returns cached result if fresh.",
    "SSRF check on URL → reject if private/reserved.",
    "ML model extracts 87 features from URL + page → returns label + confidence.",
    "HTML scraper fetches page → extracts 9 structural features.",
    "VirusTotal, AbuseIPDB, IPStack called in parallel (async).",
    "Decision engine scores all 5 signals → computes total (0–100).",
    "Label assigned: phishing / suspicious / benign.",
    "Result cached and returned to UI.",
    "UI renders gauge, charts, reasons, and expandable details.",
    "(Optional) n8n webhook triggers → logs alert + Slack notification.",
]
for i, s in enumerate(steps, 1):
    bullet(f"{i}. {s}")

# ── 9. Limitations & Future Improvements ─────────────────────────────────────
heading("9. Limitations & Future Improvements")
add_table(
    ["Limitation", "Suggested Improvement"],
    [
        ["In-memory alerts log (lost on restart)", "Persist alerts to a database (SQLite/PostgreSQL)"],
        ["No authentication on API endpoints", "Add API key or OAuth2 middleware"],
        ["Model trained on static dataset", "Implement periodic retraining pipeline"],
        ["Single-instance cache (not distributed)", "Use Redis for shared cache across instances"],
        ["IPStack free tier lacks full threat data", "Upgrade to paid tier or replace with MaxMind"],
        ["No rate limiting on /predict", "Add slowapi or similar rate limiter"],
    ],
)

# ── 10. Project Structure ─────────────────────────────────────────────────────
heading("10. Project Structure")
code_block(
    "CyberSecurity_01/\n"
    "├── api/\n"
    "│   └── main.py              # FastAPI app, endpoints, cache, SSRF guard\n"
    "├── decision/\n"
    "│   └── engine.py            # Scoring logic, label thresholds\n"
    "├── ml_model/\n"
    "│   ├── train.py             # XGBoost training + cross-validation\n"
    "│   ├── predict.py           # Feature extraction (87 features) + inference\n"
    "│   ├── phishing_model.pkl   # Trained model artifact\n"
    "│   └── dataset.csv          # Training dataset\n"
    "├── scraper/\n"
    "│   └── html_scraper.py      # BeautifulSoup HTML feature extractor\n"
    "├── threat_intel/\n"
    "│   └── api_checker.py       # VirusTotal, AbuseIPDB, IPStack async clients\n"
    "├── ui/\n"
    "│   └── app.py               # Streamlit dashboard with Plotly charts\n"
    "├── n8n/\n"
    "│   └── phishing_detection_workflow.json  # n8n automation workflow\n"
    "├── .env                     # API keys (not committed)\n"
    "├── render.yaml              # Render cloud deployment config\n"
    "└── requirements.txt         # Python dependencies"
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("PROJECT_REPORT.docx")
print("Done: PROJECT_REPORT.docx")
